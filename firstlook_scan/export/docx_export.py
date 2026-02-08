from __future__ import annotations

from io import BytesIO
from datetime import datetime
from docx import Document

from firstlook_scan.types import ScanResult
from .csv_export import scanresult_to_findings_df

def export_summary_docx_bytes(scan_result: ScanResult) -> bytes:
    doc = Document()
    doc.add_heading("FirstLook Scan — Deal Triage Summary", level=1)

    doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
    doc.add_paragraph("Note: This report is informational only and not legal advice.")

    summary = scan_result.batch_summary

    doc.add_heading("Batch Overview", level=2)
    doc.add_paragraph(f"Files scanned: {getattr(summary, 'total_files', 0)}")
    doc.add_paragraph(f"OK files: {getattr(summary, 'ok_files', 0)}")
    doc.add_paragraph(f"Needs OCR: {getattr(summary, 'needs_ocr_files', 0)}")
    doc.add_paragraph(f"Errors: {getattr(summary, 'error_files', 0)}")

    doc.add_heading("Counts by Risk Level", level=2)
    counts_risk = getattr(summary, "counts_by_risk_level", {}) or {}
    for k in ["High", "Medium", "Low"]:
        doc.add_paragraph(f"{k}: {counts_risk.get(k, 0)}", style="List Bullet")

    doc.add_heading("Counts by Clause Type", level=2)
    counts_type = getattr(summary, "counts_by_clause_type", {}) or {}
    for ct, v in counts_type.items():
        doc.add_paragraph(f"{ct}: {v}", style="List Bullet")

    df = scanresult_to_findings_df(scan_result)

    doc.add_heading("Top Findings (Risk then Confidence)", level=2)
    if df.empty:
        doc.add_paragraph("No findings available.")
    else:
        tmp = df.copy()
        tmp["risk_rank"] = tmp["risk_level"].map({"High": 3, "Medium": 2, "Low": 1}).fillna(0).astype(int)
        tmp = tmp.sort_values(["risk_rank", "confidence"], ascending=[False, False]).drop(columns=["risk_rank"])
        top = tmp.head(10)

        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "File"
        hdr[1].text = "Clause"
        hdr[2].text = "Risk"
        hdr[3].text = "Confidence"

        for _, r in top.iterrows():
            row = table.add_row().cells
            row[0].text = str(r.get("file_name", ""))
            row[1].text = str(r.get("clause_type", ""))
            row[2].text = str(r.get("risk_level", ""))
            row[3].text = f"{float(r.get('confidence', 0.0)):.2f}"

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
